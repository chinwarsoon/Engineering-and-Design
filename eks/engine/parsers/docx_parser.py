"""
DOCX Parser for EKS - Uses python-docx for Word document extraction.
"""
from docx import Document
from pathlib import Path
from typing import Any, Dict, List
from .base_parser import BaseParser

class DOCXParser(BaseParser):
    """
    Parses DOCX files to extract paragraphs and metadata.
    """
    def parse(self) -> List[Dict[str, Any]]:
        blocks = []
        doc = Document(str(self.file_path))
        
        # Combine all paragraphs into one content block for simplicity in Phase 1
        # In later phases, we might want to split by sections or headings.
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        if full_text:
            blocks.append({
                "content": "\n".join(full_text),
                "type": "text",
                "metadata": {
                    "paragraph_count": len(full_text)
                }
            })
        
        # Also extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            
            if table_data:
                content = "\n".join(["\t".join(row) for row in table_data])
                blocks.append({
                    "content": content,
                    "type": "table",
                    "metadata": {
                        "table_index": i + 1,
                        "row_count": len(table_data),
                        "column_count": len(table_data[0]) if table_data else 0
                    }
                })
                
        return blocks

    def extract_metadata(self) -> Dict[str, Any]:
        doc = Document(str(self.file_path))
        props = doc.core_properties
        return {
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "revision": props.revision,
            "last_modified_by": props.last_modified_by
        }
